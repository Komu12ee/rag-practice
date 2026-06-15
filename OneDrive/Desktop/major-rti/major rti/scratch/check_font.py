import docx
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn

doc = Document()
p = doc.add_paragraph()
r = p.add_run("क्रमांक / छत्तीसगढ़ इंफोटेक प्रमोशन सोसाइटी")
r.bold = True
r.font.size = Pt(12)

# Apply font settings
r.font.name = "Nirmala UI"
rPr = r.font.element.get_or_add_rPr()
rFonts = rPr.get_or_add_rFonts()
rFonts.set(qn('w:ascii'), 'Nirmala UI')
rFonts.set(qn('w:hAnsi'), 'Nirmala UI')
rFonts.set(qn('w:cs'), 'Nirmala UI') # Complex Script font

doc.save("scratch/test_font.docx")
print("Saved successfully!")
