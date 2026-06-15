from datetime import datetime
import io
import re
import streamlit as st

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


def _apply_font(run_or_style, font_name="Nirmala UI"):
    """Apply specific font (including complex script font for Hindi) to run or style."""
    if not DOCX_AVAILABLE:
        return
    font = run_or_style.font
    font.name = font_name
    
    # Force set Word font properties for ASCII, High-Ansi and Complex Script (Hindi)
    try:
        rPr = font.element.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn('w:ascii'), font_name)
        rFonts.set(qn('w:hAnsi'), font_name)
        rFonts.set(qn('w:cs'), 'Mangal')  # Use Mangal for Hindi (complex script) on all Windows versions
        rFonts.set(qn('w:eastAsia'), font_name)
        
        # Remove theme attributes that can override the literal font names
        for attr in ['asciiTheme', 'hAnsiTheme', 'cstheme', 'eastAsiaTheme']:
            attrib_name = qn(f'w:{attr}')
            if attrib_name in rFonts.attrib:
                del rFonts.attrib[attrib_name]
    except Exception:
        pass


def generate_analysis_docx(session_state: dict) -> bytes:
    """
    Generate a professionally formatted, highly detailed DOCX report of the AI analysis.
    Returns bytes for Streamlit download_button.
    """
    
    extracted_text = session_state.get('extracted_text', session_state.get('rti_text', ''))
    routing = session_state.get('routing_result')
    effective_dept = session_state.get('effective_department', '')
    is_chips = session_state.get('is_chips_jurisdiction', False)
    
    predicted_dept = "Unknown"
    predicted_conf = "LOW"
    routing_reasoning = ""
    section_ref = ""
    
    if routing:
        predicted_dept = getattr(routing, 'primary_department', routing.get('department', 'Unknown') if isinstance(routing, dict) else 'Unknown')
        predicted_conf = getattr(routing, 'confidence_band', routing.get('confidence', 'LOW') if isinstance(routing, dict) else 'LOW')
        routing_reasoning = getattr(routing, 'reasoning', routing.get('reasoning', '') if isinstance(routing, dict) else '')
        section_ref = getattr(routing, 'section_reference', routing.get('section_reference', '') if isinstance(routing, dict) else '')

    info = session_state.get('confirmed_info', session_state.get('extracted_info'))
    ex_flags = session_state.get('exemption_flags', [])
    layer_b = session_state.get('layer_b_res')
    balance = session_state.get('balance_res')
    recom = session_state.get('final_recom')

    if not DOCX_AVAILABLE:
        # Detailed plain text fallback
        buffer = io.BytesIO()
        lines = [
            "CHiPS — RTI Intelligence System",
            "AI Analysis Report — Advisory Only",
            "==================================",
            f"Case ID: {session_state.get('case_id', 'N/A')}",
            f"Generated: {datetime.now().strftime('%d %B %Y, %H:%M')}",
            f"Language: {session_state.get('language', 'en').upper()}",
            f"OCR Confidence: {session_state.get('ocr_confidence', 1.0):.0%}",
            "--------------------------------------------------------------------------------",
            "⚠️ LEGAL ADVISORY NOTICE: This report is AI-generated and for advisory purposes only. The Public Information Officer (PIO) remains the sole statutory authority and legal decision-maker under Section 5 of the RTI Act 2005.",
            "--------------------------------------------------------------------------------",
            "",
            "1. Application Summary",
            "----------------------",
            extracted_text if extracted_text else "No RTI application text available.",
            "",
            "2. Jurisdiction & Routing",
            "-------------------------",
            f"Recommended Department: {predicted_dept}",
            f"Analysis Confidence: {predicted_conf}",
            f"Effective Target Department: {effective_dept if effective_dept else predicted_dept}",
            f"CHiPS Jurisdiction: {'Yes' if is_chips else 'No (Requires Section 6(3) Transfer)'}",
        ]
        
        if session_state.get('pio_routing_override'):
            lines.extend([
                f"PIO Routing Override: Overridden to '{session_state['pio_routing_override']}'",
                f"Override Reason: {session_state.get('routing_correction_reason', 'N/A')}"
            ])
            
        if routing_reasoning:
            lines.extend(["", "Routing Analysis:", routing_reasoning])
        if section_ref:
            lines.extend(["", f"Legal Reference: {section_ref}"])
            
        lines.extend(["", "3. Extracted Parameters", "----------------------"])
        if info:
            lines.extend([
                f"Information Type: {getattr(info, 'information_type', 'other').upper()}",
                f"Procurement/Tender Status: {getattr(info, 'procurement_status', 'none').upper()}",
                f"Contains Personal Private Data: {'Yes' if getattr(info, 'personal_data', False) else 'No'}",
                f"Section 8(2) Override Flag: {'Yes' if getattr(info, 'public_interest_override', False) else 'No'}"
            ])
            
            entities = getattr(info, 'extracted_entities', [])
            systems = getattr(info, 'systems', [])
            if entities:
                lines.append(f"Extracted Entities: {', '.join(entities)}")
            if systems:
                lines.append(f"IT Systems / Databases: {', '.join(systems)}")
            if getattr(info, 'explanation', ''):
                lines.extend(["", "Extraction Explanation:", getattr(info, 'explanation', '')])
        else:
            lines.append("Information extraction skipped or not completed (Non-CHiPS Transfer).")
            
        lines.extend(["", "4. Exemption Analysis", "---------------------"])
        if not is_chips:
            lines.append("Exemption analysis and disclosure balancing are skipped because the target department is outside of CHiPS jurisdiction. Under Section 6(3) of the RTI Act 2005, the application must be transferred to the concerned department.")
        else:
            if not ex_flags:
                lines.append("🟢 No exemption triggers detected.")
            else:
                lines.append("🔴 Triggered Section 8(1) / 11 Exemptions:")
                for flag in ex_flags:
                    o_text = f" [OVERRIDDEN: {getattr(flag, 'override_reason', '')}]" if getattr(flag, 'is_overridden', False) else ""
                    lines.append(f"• {flag.section} - {flag.title}: {flag.reasoning} (Suggested Action: {flag.suggested_action}){o_text}")
            
            if layer_b and getattr(layer_b, 'exemptions_analysis', None):
                lines.extend(["", "Statutory RAG Analysis & Legal Citations:"])
                for ea in layer_b.exemptions_analysis:
                    status_str = "APPLICABLE (Withhold)" if ea.is_applicable else "NOT APPLICABLE (Disclose)"
                    lines.extend([
                        f"• {ea.section} status: {status_str}",
                        f"  Legal Reasoning: {ea.legal_reasoning}"
                    ])
                    if ea.exact_quotes:
                        lines.append(f"  Quotes Cited: \"{'; '.join(ea.exact_quotes)}\"")
                        
            if balance:
                lines.extend([
                    "", "Adversarial Balancing (Public Interest Test):",
                    f"  Argument for Disclosure: {getattr(balance, 'pro_disclosure_argument', 'N/A')}",
                    f"  Argument for Exemption: {getattr(balance, 'pro_exemption_argument', 'N/A')}",
                    f"  Balancing Factors: {getattr(balance, 'balancing_factors', 'N/A')}"
                ])
                
        lines.extend(["", "5. Final AI Recommendation", "--------------------------"])
        if recom:
            lines.extend([
                f"AI Synthesized Action: {getattr(recom, 'recommendation', 'N/A')}",
                f"Recommendation Confidence: {getattr(recom, 'confidence_band', 'LOW')}",
                f"Primary Legal Reasoning: {getattr(recom, 'primary_reasoning', 'N/A')}",
                f"Statutory Citations: {', '.join(getattr(recom, 'sections_applied', getattr(recom, 'statutory_citations_applied', [])))}",
                f"Suggested PIO Directive: {getattr(recom, 'suggested_pio_action', 'N/A')}"
            ])
            if getattr(recom, 'rejection_risk', None):
                lines.append(f"Refusal Appeal Risk: {recom.rejection_risk}")
            if getattr(recom, 'disclosure_risk', None):
                lines.append(f"Disclosure Security/Privacy Risk: {recom.disclosure_risk}")
        else:
            lines.append("Final recommendation details not available.")
            
        lines.extend(["", "6. PIO Notes and Decision", "-------------------------",
                      f"Final Logged Decision: {session_state.get('final_decision_value', 'Pending')}",
                      session_state.get('pio_decision_text', 'Not yet entered.')])
        
        if session_state.get('logged_audit_id'):
            lines.append(f"Audit Trail Hash Link ID: {session_state.get('logged_audit_id')}")
            
        buffer.write(("\n".join(lines)).encode('utf-8'))
        return buffer.getvalue()

    doc = Document()

    # Page setup - premium margin
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Document Header Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = title_p.add_run("CHiPS — RTI Intelligence System")
    run_title.bold = True
    run_title.font.size = Pt(18)
    run_title.font.color.rgb = RGBColor(30, 58, 138) # Dark Blue #1E3A8A
    
    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = subtitle_p.add_run("AI Analysis & PIO Decision Report — Advisory Only")
    run_sub.italic = True
    run_sub.font.size = Pt(11)
    run_sub.font.color.rgb = RGBColor(75, 85, 99) # Gray #4B5563

    # Metadata Block
    meta_p = doc.add_paragraph()
    meta_p.add_run("Case ID: ").bold = True
    meta_p.add_run(f"{session_state.get('case_id', 'N/A')}\n")
    meta_p.add_run("Generated Date: ").bold = True
    meta_p.add_run(f"{datetime.now().strftime('%d %B %Y, %H:%M')}\n")
    meta_p.add_run("Language: ").bold = True
    meta_p.add_run(f"{session_state.get('language', 'en').upper()}\n")
    meta_p.add_run("OCR Confidence: ").bold = True
    meta_p.add_run(f"{session_state.get('ocr_confidence', 1.0):.0%}\n")
    
    warning_p = doc.add_paragraph()
    w_run = warning_p.add_run("⚠️ LEGAL ADVISORY NOTICE: This report is AI-generated and for advisory purposes only. The Public Information Officer (PIO) remains the sole statutory authority and legal decision-maker under Section 5 of the RTI Act 2005.")
    w_run.font.size = Pt(9.5)
    w_run.font.color.rgb = RGBColor(220, 38, 38) # Red #DC2626
    w_run.italic = True

    # Horizontal Divider Line
    divider = doc.add_paragraph()
    divider.add_run("─" * 70).font.color.rgb = RGBColor(209, 213, 219)

    # 1. Application Summary
    h1 = doc.add_heading(level=1)
    h1_run = h1.add_run("1. Application Summary")
    h1_run.font.color.rgb = RGBColor(30, 58, 138)
    if extracted_text:
        doc.add_paragraph(extracted_text)
    else:
        doc.add_paragraph().add_run("No RTI application text available.").italic = True
        
    parse_result = session_state.get('parse_result')
    if parse_result and getattr(parse_result, 'warning', None):
        warn_p = doc.add_paragraph()
        w_run = warn_p.add_run(f"Parsing Warning: {parse_result.warning}")
        w_run.font.color.rgb = RGBColor(217, 119, 6) # Amber
        w_run.italic = True

    # 2. Jurisdiction & Routing
    h2 = doc.add_heading(level=1)
    h2_run = h2.add_run("2. Jurisdiction & Routing")
    h2_run.font.color.rgb = RGBColor(30, 58, 138)
    
    p2 = doc.add_paragraph()
    p2.add_run("Recommended Department: ").bold = True
    p2.add_run(f"{predicted_dept}\n")
    p2.add_run("Analysis Confidence: ").bold = True
    p2.add_run(f"{predicted_conf}\n")
    p2.add_run("Effective Target Department: ").bold = True
    p2.add_run(f"{effective_dept if effective_dept else predicted_dept}\n")
    p2.add_run("CHiPS Jurisdiction: ").bold = True
    p2.add_run(f"{'Yes' if is_chips else 'No (Requires Section 6(3) Transfer)'}\n")
    
    if session_state.get('pio_routing_override'):
        p2.add_run("PIO Routing Override: ").bold = True
        p2.add_run(f"Overridden to '{session_state['pio_routing_override']}'\n")
        p2.add_run("Override Reason: ").bold = True
        p2.add_run(f"{session_state.get('routing_correction_reason', 'N/A')}\n")
        
    if routing_reasoning:
        doc.add_paragraph().add_run("Routing Analysis:").bold = True
        doc.add_paragraph(routing_reasoning)
    if section_ref:
        doc.add_paragraph().add_run(f"Legal Reference: {section_ref}").italic = True

    # 3. Extracted Parameters
    h3 = doc.add_heading(level=1)
    h3_run = h3.add_run("3. Extracted Parameters")
    h3_run.font.color.rgb = RGBColor(30, 58, 138)
    if info:
        info_type = getattr(info, 'information_type', 'other')
        entities = getattr(info, 'extracted_entities', [])
        systems = getattr(info, 'systems', [])
        procurement = getattr(info, 'procurement_status', 'none')
        personal = getattr(info, 'personal_data', False)
        override = getattr(info, 'public_interest_override', False)
        explanation = getattr(info, 'explanation', '')
        
        p3 = doc.add_paragraph()
        p3.add_run("Information Type: ").bold = True
        p3.add_run(f"{info_type.upper()}\n")
        p3.add_run("Procurement/Tender Status: ").bold = True
        p3.add_run(f"{procurement.upper()}\n")
        p3.add_run("Contains Personal Private Data: ").bold = True
        p3.add_run(f"{'Yes' if personal else 'No'}\n")
        p3.add_run("Section 8(2) Override Flag (Corruption/HR Allegations): ").bold = True
        p3.add_run(f"{'Yes' if override else 'No'}\n")
        
        if entities:
            p3.add_run("Extracted Entities: ").bold = True
            p3.add_run(f"{', '.join(entities)}\n")
        if systems:
            p3.add_run("IT Systems / Databases: ").bold = True
            p3.add_run(f"{', '.join(systems)}\n")
            
        if explanation:
            doc.add_paragraph().add_run("Extraction Explanation:").bold = True
            doc.add_paragraph(explanation)
    else:
        doc.add_paragraph().add_run("Information extraction skipped or not completed (Non-CHiPS Transfer).").italic = True

    # 4. Exemption Analysis
    h4 = doc.add_heading(level=1)
    h4_run = h4.add_run("4. Exemption Analysis")
    h4_run.font.color.rgb = RGBColor(30, 58, 138)
    if not is_chips:
        doc.add_paragraph("Exemption analysis and disclosure balancing are skipped because the target department is outside of CHiPS jurisdiction. Under Section 6(3) of the RTI Act 2005, the application must be transferred to the concerned department.")
    else:
        if not ex_flags:
            doc.add_paragraph("🟢 No exemption triggers detected. The requested information does not appear to trigger any Section 8(1) exemption clauses and falls under standard disclosure guidelines.")
        else:
            doc.add_paragraph("🔴 Triggered Section 8(1) / 11 Exemptions:").bold = True
            for flag in ex_flags:
                fp = doc.add_paragraph(style='List Bullet')
                r_title = fp.add_run(f"{flag.section} - {flag.title}: ")
                r_title.bold = True
                fp.add_run(f"{flag.reasoning} (Suggested Action: {flag.suggested_action})")
                if getattr(flag, 'is_overridden', False):
                    o_run = fp.add_run(f" [OVERRIDDEN: {getattr(flag, 'override_reason', '')}]")
                    o_run.italic = True
                    o_run.font.color.rgb = RGBColor(217, 119, 6)

        if layer_b and getattr(layer_b, 'exemptions_analysis', None):
            doc.add_paragraph().add_run("Statutory RAG Analysis & Legal Citations:").bold = True
            for ea in layer_b.exemptions_analysis:
                eap = doc.add_paragraph()
                status_str = "APPLICABLE (Withhold)" if ea.is_applicable else "NOT APPLICABLE (Disclose)"
                eap.add_run(f"• {ea.section} status: ").bold = True
                eap.add_run(f"{status_str}\n")
                eap.add_run(f"Legal Reasoning: {ea.legal_reasoning}\n")
                if ea.exact_quotes:
                    eap.add_run("Quotes Cited: ").italic = True
                    eap.add_run(f"\"{'; '.join(ea.exact_quotes)}\"")

            if getattr(layer_b, 'overall_explanation', None):
                doc.add_paragraph().add_run("Overall Exemption Summary:").bold = True
                doc.add_paragraph(layer_b.overall_explanation)

        if balance:
            doc.add_paragraph().add_run("Adversarial Balancing (Public Interest Test):").bold = True
            bp = doc.add_paragraph()
            bp.add_run("Argument for Disclosure: ").bold = True
            bp.add_run(f"{getattr(balance, 'pro_disclosure_argument', 'N/A')}\n\n")
            bp.add_run("Argument for Exemption: ").bold = True
            bp.add_run(f"{getattr(balance, 'pro_exemption_argument', 'N/A')}\n\n")
            bp.add_run("Balancing Factors: ").bold = True
            bp.add_run(f"{getattr(balance, 'balancing_factors', 'N/A')}")

    # 5. Final Recommendation
    h5 = doc.add_heading(level=1)
    h5_run = h5.add_run("5. Final AI Recommendation")
    h5_run.font.color.rgb = RGBColor(30, 58, 138)
    if recom:
        p5 = doc.add_paragraph()
        p5.add_run("AI Synthesized Action: ").bold = True
        r_run = p5.add_run(f"{getattr(recom, 'recommendation', 'N/A')}\n")
        r_run.bold = True
        p5.add_run("Recommendation Confidence: ").bold = True
        p5.add_run(f"{getattr(recom, 'confidence_band', 'LOW')}\n")
        
        p5.add_run("Primary Legal Reasoning:\n")
        p5.add_run(f"{getattr(recom, 'primary_reasoning', 'N/A')}\n\n")
        
        p5.add_run("Statutory Citations: ").bold = True
        p5.add_run(f"{', '.join(getattr(recom, 'sections_applied', getattr(recom, 'statutory_citations_applied', [])))}\n")
        p5.add_run("Suggested PIO Directive: ").bold = True
        p5.add_run(f"{getattr(recom, 'suggested_pio_action', 'N/A')}\n")
        
        if getattr(recom, 'rejection_risk', None):
            p5.add_run("Refusal Appeal Risk: ").bold = True
            p5.add_run(f"{recom.rejection_risk}\n")
        if getattr(recom, 'disclosure_risk', None):
            p5.add_run("Disclosure Security/Privacy Risk: ").bold = True
            p5.add_run(f"{recom.disclosure_risk}\n")
    else:
        doc.add_paragraph("Final recommendation details not available.")

    # 6. PIO Notes and Decision
    h6 = doc.add_heading(level=1)
    h6_run = h6.add_run("6. PIO Notes and Decision")
    h6_run.font.color.rgb = RGBColor(30, 58, 138)
    doc.add_paragraph().add_run(f"Final Logged Decision: {session_state.get('final_decision_value', 'Pending')}").bold = True
    doc.add_paragraph(session_state.get('pio_decision_text', 'Not yet entered.'))
    
    audit_id = session_state.get('logged_audit_id')
    if audit_id:
        doc.add_paragraph().add_run(f"Audit Trail Hash Link ID: {audit_id}").italic = True

    # Footer line
    doc.add_paragraph('─' * 80)
    footer_p = doc.add_paragraph('This document is intended for internal government use only. AI-Generated Draft.')
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Final pass to apply font formatting and clear theme overrides on all runs in the document
    if DOCX_AVAILABLE:
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                _apply_font(run, 'Nirmala UI')
                
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            _apply_font(run, 'Nirmala UI')
                            
        for section in doc.sections:
            for header_p in section.header.paragraphs:
                for run in header_p.runs:
                    _apply_font(run, 'Nirmala UI')
            for footer_p in section.footer.paragraphs:
                for run in footer_p.runs:
                    _apply_font(run, 'Nirmala UI')

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def render_export_buttons(session_state: dict):
    """Render download buttons in Streamlit. Call at end of Step 3."""
    st.divider()
    st.subheader('📥 Export Analysis')
    
    if not DOCX_AVAILABLE:
        st.warning("⚠️ python-docx library is not installed. Exports will be formatted as plain text files.")

    col1, col2 = st.columns(2)
    with col1:
        docx_bytes = generate_analysis_docx(session_state)
        ext = 'docx' if DOCX_AVAILABLE else 'txt'
        mime_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document' if DOCX_AVAILABLE else 'text/plain'
        st.download_button(
            label=f'⬇️ Download AI Analysis as {ext.upper()}',
            data=docx_bytes,
            file_name=f"RTI_Analysis_{session_state.get('case_id', 'report').replace('/', '_')}.{ext}",
            mime=mime_type
        )
    with col2:
        st.info('Export tip: Open the DOCX and Print → Save as PDF to generate a PDF report.')
