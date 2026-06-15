from datetime import datetime
import io
import re
import os
import logging
import sys
from pathlib import Path
import streamlit as st

from sarvam_client import call_sarvam_chat, SARVAM_API_KEY

logger = logging.getLogger(__name__)

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

DECISION_PATHS = {
    'Approve — Full Disclosure': 'APPROVE',
    'Partially Approve — Partial Disclosure': 'PARTIAL',
    'Reject — Exempt': 'REJECT',
    'Transfer — Section 6(3)': 'TRANSFER',
    'Section 5(4) — Information Collection': 'SECTION54',
}


def generate_structured_legal_response(
    analysis_result,
    appellant_name: str = None,
    rti_date: str = None,
    rti_subject: str = None,
    reply_date: str = None,
):
    """
    Generate citation-aware response outputs from an RTIAnalysisResult.

    This extends the existing response generation module without changing the
    existing DOCX/letter generation path used by generate_response_letter().
    """
    project_root = Path(__file__).resolve().parent.parent
    src_dir = project_root / "src"
    if str(src_dir) not in sys.path:
        sys.path.insert(0, str(src_dir))

    from generation.response_templates import ResponseGenerator

    return ResponseGenerator().generate_all(
        analysis=analysis_result,
        appellant_name=appellant_name,
        rti_date=rti_date,
        rti_subject=rti_subject,
        reply_date=reply_date,
    )


def detect_language(text: str) -> str:
    """Detect language based on Hindi vs English character ratio."""
    if not text or not text.strip():
        return "english"
    
    hindi_chars = sum(1 for c in text if '\u0900' <= c <= '\u097F')
    english_chars = sum(1 for c in text if c.isalpha())
    
    if english_chars == 0:
        return "hindi" if hindi_chars > 0 else "english"
        
    ratio = hindi_chars / english_chars
    if ratio > 0.3:
        return "hindi"
    elif ratio > 0.05:
        return "mixed"
    return "english"


def _clean_json_response(raw: str) -> str:
    """Extract and clean outer JSON string from LLM responses using bracket matching."""
    start = raw.find('{')
    if start == -1:
        return raw
    
    depth = 0
    in_string = False
    escape = False
    end_idx = -1
    
    for i in range(start, len(raw)):
        char = raw[i]
        if in_string:
            if escape:
                escape = False
            elif char == '\\':
                escape = True
            elif char == '"':
                in_string = False
        else:
            if char == '"':
                in_string = True
            elif char == '{':
                depth += 1
            elif char == '}':
                depth -= 1
                if depth == 0:
                    end_idx = i
                    break
                    
    if end_idx != -1:
        raw = raw[start:end_idx+1]
        
    raw = re.sub(r"```json\s*", "", raw, flags=re.IGNORECASE)
    raw = re.sub(r"```\s*$", "", raw)
    raw = raw.strip()
    raw = re.sub(r",\s*([\]}])", r"\1", raw)
    return raw


def _apply_font(run_or_style, font_name="Nirmala UI"):
    """Apply specific font (including complex script font for Hindi) to run or style."""
    if not DOCX_AVAILABLE:
        return
    font = run_or_style.font
    font.name = font_name
    
    # Force set Word font properties for ASCII, High-Ansi and Complex Script (Hindi)
    try:
        rPr = font.element.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn('w:ascii'), font_name)
        rFonts.set(qn('w:hAnsi'), font_name)
        rFonts.set(qn('w:cs'), 'Mangal')  # Use Mangal for Hindi (complex script) on all Windows versions
        rFonts.set(qn('w:eastAsia'), font_name)
        
        # Remove theme attributes that can override the literal font names
        for attr in ['asciiTheme', 'hAnsiTheme', 'cstheme', 'eastAsiaTheme']:
            attrib_name = qn(f'w:{attr}')
            if attrib_name in rFonts.attrib:
                del rFonts.attrib[attrib_name]
    except Exception as e:
        logger.warning(f"Could not apply XML font elements: {e}")


def _extract_applicant_details_fallback(text: str, session_state: dict) -> tuple:
    """Extract applicant name and address using regex and confirmed entities (fallback)."""
    applicant_name = 'The Applicant / आवेदक'
    applicant_address = 'Address not specified / पता उल्लेखित नहीं'
    
    # Try regex match for Name
    name_match = re.search(r'(?:Name|नाम|आवेदक का नाम|Applicant Name)\s*[:ः-]\s*([^\n\r]+)', text, re.IGNORECASE)
    if name_match:
        applicant_name = name_match.group(1).strip()
    else:
        info = session_state.get('confirmed_info', session_state.get('extracted_info'))
        if info:
            for ent in getattr(info, 'extracted_entities', []):
                if any(title in ent.lower() for title in ["shri", "smt", "kumar", "singh", "devi", "श्रीमती", "श्री"]):
                    applicant_name = ent
                    break

    # Try regex match for Address
    address_match = re.search(r'(?:Address|पता|आवेदक का पता|Correspondence Address)\s*[:ः-]\s*([^\n\r]+(?:\n\s*[^\n\r]+){0,2})', text, re.IGNORECASE)
    if address_match:
        applicant_address = address_match.group(1).strip().replace('\n', ', ')
        
    return applicant_name, applicant_address


def _generate_letter_content(
    text: str,
    decision_path: str,
    language: str,
    effective_dept: str,
    active_exemptions: list,
    entities: list,
    systems: list,
    info_type: str,
    pio_notes: str
) -> dict:
    """Generate the response letter details (summary, body, applicant info) in a single LLM call."""
    if not SARVAM_API_KEY:
        return {}

    decision_mapping = {
        'APPROVE': 'Approve / Full Disclosure (CHiPS holds records)',
        'PARTIAL': 'Partially Approve / Partial Disclosure (Exemptions apply under Section 8(1) and redacted under Section 10)',
        'REJECT': 'Reject / Exempt from Disclosure',
        'TRANSFER': f'Transfer to concerned department ({effective_dept}) under Section 6(3)',
        'SECTION54': 'Information Collection from other division under Section 5(4)',
    }
    
    decision_text = decision_mapping.get(decision_path, decision_path)
    exemptions_str = ", ".join(active_exemptions) if active_exemptions else "None"
    
    prompt = f"""You are an expert legal information extraction assistant and Public Information Officer (PIO) drafting assistant for Chhattisgarh Infotech Promotion Society (CHiPS).
Your task is to analyze the provided RTI application text and draft an official response letter, including metadata extraction.

─────────────────────────────────────────────
RTI APPLICATION ORIGINAL TEXT
─────────────────────────────────────────────
{text[:10000]}

─────────────────────────────────────────────
PIO-CONFIRMED METADATA & DECISIONS
─────────────────────────────────────────────
Target Language: {language}
Final Decision Path: {decision_text}
Confirmed Target Department: {effective_dept}
Active Legal Exemption Sections: {exemptions_str}
Confirmed Key Entities: {', '.join(entities) if entities else 'None'}
Confirmed IT Systems: {', '.join(systems) if systems else 'None'}
Information Category: {info_type}
PIO Custom Notes & Comments: {pio_notes}

─────────────────────────────────────────────
OUTPUT FORMAT
─────────────────────────────────────────────
Return ONLY valid JSON matching this schema:
{{
  "applicant_name": "Extract the applicant's name. Default to 'The Applicant / आवेदक' if not found.",
  "applicant_address": "Extract the applicant's address. Default to 'Address not specified / पता उल्लेखित नहीं' if not found.",
  "rti_summary": "A concise subject-matter noun phrase summarizing what the request is about in 10-15 words (in the target language, e.g. 'deployment of Network Engineers under CG SWAN project' or 'procurement of State Data Centre servers')",
  "letter_body": "The drafted body paragraphs of the response letter (in the target language)."
}}

INSTRUCTIONS FOR LETTER BODY:
1. Write in a formal, official Government of Chhattisgarh memorandum style.
2. The tone must be polite, professional, and legally precise.
3. DO NOT use generic placeholder language (like 'With reference to the requested information...'). Instead, refer explicitly to the actual subject matter.
4. Explain the decision clearly based on the PIO-confirmed decisions:
   - For APPROVE: State that the request was examined, belongs to CHiPS, is disclosable, and the requested records are enclosed.
   - For PARTIAL: Explain that the request was examined. Some parts are exempt under {exemptions_str} and have been redacted/withheld under Section 10 (severability), while the remaining records are enclosed.
   - For REJECT: State that the request was examined and the information is exempt from disclosure under Section {exemptions_str} of the RTI Act, 2005. Explain the legal basis clearly.
   - For TRANSFER: State that the requested information does not pertain to CHiPS and is maintained by "{effective_dept}". Under Section 6(3) of the Act, the application is being transferred to the concerned PIO of "{effective_dept}" within the statutory 5-day limit.
   - For SECTION54: State that assistance from the concerned officer (deemed PIO) is being sought under Section 5(4) to collect the requested information.
5. Incorporate the PIO's Custom Notes: "{pio_notes}" naturally into the explanation.
6. Draft the body paragraphs ONLY. Do not write the header (e.g. Logo, CHiPS, Memo number, From, To, Date, Subject, Ref) and do not write the signature block or disclaimer footer. Python will fill those automatically.
7. Write both 'rti_summary' and 'letter_body' in the requested language ({language}). If language is mixed, write in Hindi using common English terms in Devenagari script where appropriate.

Ensure all keys and string values are properly enclosed in double quotes.
Ensure there is absolutely NO conversational preamble, NO thought process, NO explanation, NO reconsiderations, and NO other text before or after the JSON block. Start directly with '{{' and end with '}}'. Do not output multiple JSON objects, do not reconsider your results, and do not repeat yourself.
"""

    try:
        content = call_sarvam_chat(
            messages=[{"role": "user", "content": prompt}],
            temperature=0.1
        )
        cleaned_content = _clean_json_response(content)
        import json
        parsed = json.loads(cleaned_content)
        return parsed
    except Exception as e:
        logger.error(f"Failed to generate response letter content: {e}")
        return {}


def _generate_fallback_content(
    decision_path: str,
    language: str,
    effective_dept: str,
    active_exemptions: list,
    pio_notes: str
) -> dict:
    """Generate professional subject-aware fallback text if LLM call fails."""
    summary_val = "the requested information"
    sections_str = ", ".join(active_exemptions) if active_exemptions else "Section 8(1)"
    
    if language == 'hindi' or language == 'mixed':
        summary_val = "वांछित जानकारी"
        if decision_path == 'APPROVE':
            body = (
                f"विषयांतर्गत आपके आवेदन पत्र के माध्यम से चाही गई जानकारी के संबंध में "
                "चिप्स कार्यालय द्वारा परीक्षण किया गया। उक्त जानकारी चिप्स के कार्यक्षेत्र के अंतर्गत आती है "
                "तथा प्रकटीकरण योग्य है। अतः वांछित जानकारी/दस्तावेज संलग्न कर प्रेषित किए जा रहे हैं।\n"
            )
        elif decision_path == 'PARTIAL':
            body = (
                f"विषयांतर्गत आपके आवेदन पत्र का परीक्षण किया गया। आपके द्वारा चाही गई जानकारी में से "
                f"कुछ अंश व्यक्तिगत/तृतीय पक्षीय गोपनीयता से संबंधित होने के कारण अधिनियम की धारा {sections_str} "
                "के अंतर्गत प्रकटीकरण से छूट प्राप्त है। अतः उक्त छूट प्राप्त अंशों को धारा 10 (पृथक्करणीयता) के "
                "तहत पृथक कर रोक दिया गया है, तथा शेष जानकारी संलग्न कर प्रेषित की जा रही है।\n"
            )
        elif decision_path == 'REJECT':
            body = (
                f"विषयांतर्गत आपके आवेदन पत्र का परीक्षण किया गया। आपके द्वारा चाही गई जानकारी "
                f"सूचना का अधिकार अधिनियम, 2005 की धारा {sections_str} के अपवादों के अंतर्गत प्रकटीकरण से पूर्णतः "
                "छूट प्राप्त होने के कारण प्रदाय किया जाना संभव नहीं है। अतः आपका आवेदन अस्वीकृत किया जाता है।\n"
            )
        else: # TRANSFER
            body = (
                f"विषयांतर्गत आपके आवेदन पत्र का परीक्षण किया गया। आपके द्वारा चाही गई जानकारी "
                f"चिप्स के क्षेत्राधिकार के अंतर्गत नहीं आती है, बल्कि यह मुख्य रूप से \"{effective_dept}\" से "
                "संबंधित है। अतः सूचना का अधिकार अधिनियम, 2005 की धारा 6(3) के प्रावधानों के अंतर्गत आपका आवेदन "
                f"आवश्यक कार्यवाही हेतु 05 दिवस के भीतर जन सूचना अधिकारी, {effective_dept} को अंतरित किया जा रहा है।\n"
            )
        if pio_notes.strip():
            body += f"\nटिप्पणी: {pio_notes}\n"
    else: # english
        summary_val = "the requested information"
        if decision_path == 'APPROVE':
            body = (
                f"With reference to your RTI application, the information sought has been examined. "
                "The requested records fall under the jurisdiction of CHiPS and are disclosable. "
                "Accordingly, the requested documents are enclosed herewith.\n"
            )
        elif decision_path == 'PARTIAL':
            body = (
                f"With reference to your RTI application, the information sought has been examined. "
                f"Certain portions are exempt from disclosure under Section {sections_str} of the RTI Act, 2005 "
                "and have been severed under Section 10 (severability) and withheld, and the remaining records are enclosed.\n"
            )
        elif decision_path == 'REJECT':
            body = (
                f"With reference to your RTI application, the information sought has been examined. "
                f"The requested records are exempt from disclosure under Section {sections_str} of the RTI Act, 2005. "
                "Hence, the information cannot be provided and your application is rejected.\n"
            )
        else: # TRANSFER
            body = (
                f"With reference to your RTI application, the information sought does not pertain to CHiPS. "
                f"The records are maintained by \"{effective_dept}\". Accordingly, your application is being "
                f"transferred to the PIO of {effective_dept} under Section 6(3) of the RTI Act, 2005 for direct disposal.\n"
            )
        if pio_notes.strip():
            body += f"\nNote: {pio_notes}\n"
            
    return {
        "rti_summary": summary_val,
        "letter_body": body
    }


def ensure_default_templates():
    """Ensure templates directory and default template files exist."""
    project_root = Path(__file__).resolve().parent.parent
    templates_dir = project_root / "templates"
    templates_dir.mkdir(exist_ok=True)
    
    if not DOCX_AVAILABLE:
        return
        
    template_names = [
        "approve_english", "approve_hindi",
        "partial_english", "partial_hindi",
        "reject_english", "reject_hindi",
        "transfer_english", "transfer_hindi",
        "section54_english", "section54_hindi"
    ]
    
    for name in template_names:
        file_path = templates_dir / f"{name}.docx"
        if not file_path.exists():
            try:
                # Create a simple template with custom margins
                doc = Document()
                for section in doc.sections:
                    section.top_margin = Inches(1)
                    section.bottom_margin = Inches(1)
                    section.left_margin = Inches(1)
                    section.right_margin = Inches(1)
                
                if 'Normal' in doc.styles:
                    _apply_font(doc.styles['Normal'], 'Nirmala UI')
                    
                doc.add_paragraph(f"[Default Template: {name}]")
                doc.save(str(file_path))
                print(f"[response_letter.py] Created default template: {file_path}")
            except Exception as e:
                logger.warning(f"Could not create default template {name}: {e}")


def generate_response_letter(session_state: dict) -> bytes:
    """
    Generate official RTI response letter in government memo format.
    Incorporates all verified PIO overrides and metadata from session_state.
    Returns DOCX/TXT bytes.
    """
    ensure_default_templates()
    
    case_id = session_state.get('case_id', f"RTI/CHiPS/{datetime.now().year}/0001")
    extracted_text = session_state.get('extracted_text', session_state.get('rti_text', ''))
    
    # 1. Detect Language
    lang_val = session_state.get('language', session_state.get('rti_language', 'en'))
    if lang_val in ('hi', 'hindi'):
        language = 'hindi'
    elif lang_val in ('en', 'english'):
        language = 'english'
    else:
        language = detect_language(extracted_text)
    
    session_state['rti_language'] = language
    
    # 2. Pull all confirmed metadata (with PIO overrides)
    effective_dept = session_state.get('effective_department', '')
    if not effective_dept:
        routing = session_state.get('routing_result')
        if routing:
            if isinstance(routing, dict):
                effective_dept = routing.get('department_name', routing.get('department', 'Concerned Department'))
            else:
                effective_dept = getattr(routing, 'department_name', getattr(routing, 'primary_department', 'Concerned Department'))
    if not effective_dept:
        effective_dept = 'Concerned Department'
        
    decision = session_state.get('final_decision_value', '')
    path = DECISION_PATHS.get(decision, 'APPROVE')
    
    # Exemption flags
    active_exemptions = []
    for flag in session_state.get('exemption_flags', []):
        if not getattr(flag, 'is_overridden', False):
            active_exemptions.append(flag.section)
            
    if not active_exemptions and path in ('PARTIAL', 'REJECT'):
        recom = session_state.get('final_recom')
        if recom:
            active_exemptions = getattr(recom, 'sections_applied', getattr(recom, 'statutory_citations_applied', []))
        if not active_exemptions:
            active_exemptions = ['Section 8(1)']
            
    # Confirmed entities, systems, and explanations
    info = session_state.get('confirmed_info', session_state.get('extracted_info'))
    entities = []
    systems = []
    info_type = "other"
    if info:
        entities = getattr(info, 'extracted_entities', [])
        systems = getattr(info, 'systems', [])
        info_type = getattr(info, 'information_type', 'other')
        
    pio_notes = session_state.get('pio_decision_text', '')
    year = datetime.now().year
    date_str = datetime.now().strftime('%d/%m/%Y')
    
    # 3. Generate Subject & Letter Body & Extract Applicant Details (Single LLM Call)
    content_dict = _generate_letter_content(
        text=extracted_text,
        decision_path=path,
        language=language,
        effective_dept=effective_dept,
        active_exemptions=active_exemptions,
        entities=entities,
        systems=systems,
        info_type=info_type,
        pio_notes=pio_notes
    )
    
    applicant_name = content_dict.get('applicant_name')
    applicant_address = content_dict.get('applicant_address')
    
    # If LLM failed to extract name/address, use regex fallbacks
    if not applicant_name or applicant_name == 'The Applicant / आवेदक':
        h_name, h_addr = _extract_applicant_details_fallback(extracted_text, session_state)
        if not applicant_name:
            applicant_name = h_name
        if not applicant_address or applicant_address == 'Address not specified / पता उल्लेखित नहीं':
            applicant_address = h_addr

    if not content_dict or not content_dict.get('letter_body'):
        content_dict = _generate_fallback_content(path, language, effective_dept, active_exemptions, pio_notes)
        
    rti_summary = content_dict.get('rti_summary', 'the requested information')
    body_text = content_dict.get('letter_body', '')
    
    # ─── FALLBACK IF DOCX NOT AVAILABLE ──────────────────────────
    if not DOCX_AVAILABLE:
        buffer = io.BytesIO()
        text = f"""छत्तीसगढ़ इंफोटेक प्रमोशन सोसाइटी (चिप्स)
CHHATTISGARH INFOTECH PROMOTION SOCIETY (CHiPS)
(जन सूचना अधिकारी कार्यालय / PIO Office)

क्रमांक/{case_id}/जसूअ/सु.अधि./56/चिप्स/{year}/                    रायपुर, दिनांक {date_str}

प्रेषक : श्रीधर दीवान, जन सूचना अधिकारी (PIO)
प्रति : {applicant_name}, {applicant_address}

विषयः- सूचना का अधिकार अधिनियम, 2005 के तहत जानकारी प्रदाय करने के संबंध में (विषय: {rti_summary})
(Subject: Provision of information under the Right to Information Act, 2005 for: {rti_summary})

महोदय/महोदया (Dear Sir/Madam),

{body_text}
"""
        if path in ('REJECT', 'PARTIAL'):
            text += f"""
अपील की जानकारी (Appeal Information):
यदि आप उपरोक्त निर्णय से संतुष्ट नहीं हैं, तो निर्णय प्राप्ति के 30 दिवस के भीतर प्रथम अपीलीय अधिकारी, चिप्स के समक्ष प्रथम अपील प्रस्तुत कर सकते हैं।
(If you are not satisfied, you may file a First Appeal within 30 days to the First Appellate Authority, CHiPS.)
"""

        text += f"""
भवदीय (Yours faithfully),


श्रीधर दीवान (Sridhar Diwan)
जन सूचना अधिकारी / Public Information Officer
चिप्स, छत्तीसगढ़ शासन / CHiPS, Chhattisgarh Govt.

[Note: AI-Generated Draft — Review Required]
"""
        buffer.write(text.encode('utf-8'))
        return buffer.getvalue()

    # ─── DOCX GENERATION ──────────────────────────────────────────
    project_root = Path(__file__).resolve().parent.parent
    templates_dir = project_root / "templates"
    
    # Load template based on path and language key
    lang_key = 'hindi' if language in ('hindi', 'mixed') else 'english'
    template_path = templates_dir / f"{path.lower()}_{lang_key}.docx"
    
    if template_path.exists():
        try:
            doc = Document(str(template_path))
            if len(doc.paragraphs) > 0 and "[Default Template" in doc.paragraphs[0].text:
                doc.paragraphs[0].text = ""
        except Exception as e:
            logger.warning(f"Failed to load template {template_path}, falling back: {e}")
            doc = Document()
    else:
        doc = Document()

    # Setup margins and default document style font
    for s in doc.sections:
        s.top_margin = Inches(1)
        s.bottom_margin = Inches(1)
        s.left_margin = Inches(1)
        s.right_margin = Inches(1)

    if 'Normal' in doc.styles:
        _apply_font(doc.styles['Normal'], 'Nirmala UI')

    # 1. Logo
    logo_path = project_root / "assets" / "chips_logo.png"
    if logo_path.exists():
        try:
            doc.add_picture(str(logo_path), width=Inches(1.2))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e:
            logger.warning(f"Failed to insert logo: {e}")

    # 2. Header
    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    r1 = h.add_run('छत्तीसगढ़ इंफोटेक प्रमोशन सोसाइटी (चिप्स)\n')
    r1.bold = True
    r1.font.size = Pt(12)
    r1.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
    _apply_font(r1, 'Nirmala UI')
    
    r2 = h.add_run('CHHATTISGARH INFOTECH PROMOTION SOCIETY (CHiPS)\n')
    r2.bold = True
    r2.font.size = Pt(12)
    r2.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
    _apply_font(r2, 'Nirmala UI')
    
    r3 = h.add_run('लोक सूचना अधिकारी कार्यालय / Office of the Public Information Officer')
    r3.italic = True
    r3.font.size = Pt(10)
    _apply_font(r3, 'Nirmala UI')
    
    p_div = doc.add_paragraph()
    r_div = p_div.add_run('─' * 60)
    _apply_font(r_div, 'Nirmala UI')
    p_div.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 3. File Number & Date
    p_meta = doc.add_paragraph()
    r_meta_no = p_meta.add_run(f'क्रमांक/{case_id}/जसूअ/सु.अधि./56/चिप्स/{year}/')
    r_meta_no.bold = True
    _apply_font(r_meta_no, 'Nirmala UI')
    
    r_meta_date = p_meta.add_run(f'\t\tरायपुर, दिनांक: {date_str}')
    _apply_font(r_meta_date, 'Nirmala UI')

    # 4. Sender Info
    p_sender = doc.add_paragraph()
    r_send_lbl = p_sender.add_run('प्रेषक : ')
    r_send_lbl.bold = True
    _apply_font(r_send_lbl, 'Nirmala UI')
    
    r_send_val = p_sender.add_run('श्रीधर दीवान, जन सूचना अधिकारी (PIO), चिप्स')
    _apply_font(r_send_val, 'Nirmala UI')

    # 5. Recipient Info
    p_rec = doc.add_paragraph()
    r_rec_lbl = p_rec.add_run('प्रति : ')
    r_rec_lbl.bold = True
    _apply_font(r_rec_lbl, 'Nirmala UI')
    
    r_rec_val = p_rec.add_run(f'{applicant_name},\nपता – {applicant_address}')
    _apply_font(r_rec_val, 'Nirmala UI')

    doc.add_paragraph('')

    # 6. Subject Line
    p_sub = doc.add_paragraph()
    r_sub_lbl = p_sub.add_run('विषयः- ')
    r_sub_lbl.bold = True
    _apply_font(r_sub_lbl, 'Nirmala UI')
    
    r_sub_val = p_sub.add_run(f'सूचना का अधिकार अधिनियम, 2005 के तहत जानकारी प्रदाय करने के संबंध में (विषय: {rti_summary}).\n')
    _apply_font(r_sub_val, 'Nirmala UI')
    
    r_sub_en = p_sub.add_run(f'(Subject: Regarding provision of information under the Right to Information Act, 2005 for: {rti_summary}.)')
    r_sub_en.italic = True
    _apply_font(r_sub_en, 'Nirmala UI')

    doc.add_paragraph('')

    # 7. Salutation
    p_sal = doc.add_paragraph()
    r_sal = p_sal.add_run('महोदय/महोदया (Dear Sir/Madam),')
    _apply_font(r_sal, 'Nirmala UI')
    doc.add_paragraph('')
    
    # 8. Dynamic Body
    for paragraph in body_text.split("\n\n"):
        if paragraph.strip():
            p_body = doc.add_paragraph()
            r_body = p_body.add_run(paragraph.strip())
            _apply_font(r_body, 'Nirmala UI')
            p_body.paragraph_format.line_spacing = 1.15
            p_body.paragraph_format.space_after = Pt(6)

    # 9. Appeal Rights
    if path in ('REJECT', 'PARTIAL'):
        doc.add_paragraph('')
        doc.add_heading('अपील के अधिकार की जानकारी (First Appeal Details):', level=2)
        p_appeal = doc.add_paragraph()
        r_app_hi = p_appeal.add_run(
            'यदि आप उपरोक्त जानकारी/निर्णय से संतुष्ट नहीं हैं, तो इस पत्र की प्राप्ति के 30 दिवस के भीतर प्रथम अपीलीय अधिकारी, चिप्स के समक्ष अपील दायर कर सकते हैं।\n'
        )
        _apply_font(r_app_hi, 'Nirmala UI')
        
        r_app_en = p_appeal.add_run(
            'Address: First Appellate Authority, CHiPS, Sector-19, Nava Raipur, Atal Nagar, Raipur (C.G.)\n'
        )
        r_app_en.italic = True
        _apply_font(r_app_en, 'Nirmala UI')
        
        r_app_hi2 = p_appeal.add_run(
            'इसके पश्चात भी संतुष्ट न होने पर धारा 19(3) के तहत छत्तीसगढ़ राज्य सूचना आयोग, नवा रायपुर में द्वितीय अपील प्रस्तुत की जा सकती है।'
        )
        _apply_font(r_app_hi2, 'Nirmala UI')

    doc.add_paragraph('')
    p_close = doc.add_paragraph()
    r_close = p_close.add_run('भवदीय (Yours faithfully),')
    _apply_font(r_close, 'Nirmala UI')
    
    doc.add_paragraph('\n\n')
    
    p_line = doc.add_paragraph()
    p_line.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r_line = p_line.add_run('________________________________')
    _apply_font(r_line, 'Nirmala UI')
    
    p_sig = doc.add_paragraph()
    p_sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    r_sig_name = p_sig.add_run('**श्रीधर दीवान (Sridhar Diwan)**\n')
    r_sig_name.bold = True
    _apply_font(r_sig_name, 'Nirmala UI')
    
    r_sig_title = p_sig.add_run('जन सूचना अधिकारी / Public Information Officer\n')
    _apply_font(r_sig_title, 'Nirmala UI')
    
    r_sig_dept = p_sig.add_run('चिप्स, छत्तीसगढ़ शासन / CHiPS, Chhattisgarh Govt.')
    _apply_font(r_sig_dept, 'Nirmala UI')

    # 10. AI Disclaimer at very bottom
    doc.add_paragraph('─' * 60)
    p_disc = doc.add_paragraph()
    
    r_disc_hi = p_disc.add_run('विधिक चेतावनी (Disclaimer): ')
    r_disc_hi.bold = True
    _apply_font(r_disc_hi, 'Nirmala UI')
    
    r_disc_hi_text = p_disc.add_run(
        'यह पत्र चिप्स आर.टी.आई. इंटेलिजेंस सिस्टम (एआई) के सहयोग से तैयार किया गया है। जन सूचना अधिकारी द्वारा इसका पूर्ण अवलोकन व प्रमाणीकरण किया गया है। विधिक दायित्व केवल अधिकारी का होगा।\n'
    )
    _apply_font(r_disc_hi_text, 'Nirmala UI')
    
    r_disc_en = p_disc.add_run(
        '(This response letter was generated with AI assistance and validated by the PIO. The final decision is the sole responsibility of the PIO.)'
    )
    r_disc_en.italic = True
    _apply_font(r_disc_en, 'Nirmala UI')

    # Final pass to apply font formatting and clear theme overrides on all runs in the document
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            _apply_font(run, 'Nirmala UI')
            
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        _apply_font(run, 'Nirmala UI')
                        
    for section in doc.sections:
        for header_p in section.header.paragraphs:
            for run in header_p.runs:
                _apply_font(run, 'Nirmala UI')
        for footer_p in section.footer.paragraphs:
            for run in footer_p.runs:
                _apply_font(run, 'Nirmala UI')

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def render_response_letter_section(session_state: dict):
    """Add to bottom of Step 3 after decision is finalized."""
    if not session_state.get('decision_finalized', False):
        return

    st.divider()
    st.subheader('📄 Generate Official Response Letter')
    st.caption('Generates a formal RTI response letter in CHiPS memo format. Review before sending.')

    if st.button('📝 Generate Response Letter Draft', type='primary'):
        with st.spinner('Generating official response letter...'):
            letter_bytes = generate_response_letter(session_state)
            st.session_state['response_letter_bytes'] = letter_bytes
            st.success('Response letter generated. Download and review before use.')
            st.warning('⚠️ AI-Generated Draft — Review Required before sending to applicant.')

    if session_state.get('response_letter_bytes'):
        case_id = session_state.get('case_id', 'response')
        safe_case_id = case_id.replace('/', '_').replace('\\', '_')
        
        ext = 'docx' if DOCX_AVAILABLE else 'txt'
        mime_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document' if DOCX_AVAILABLE else 'text/plain'
        
        st.download_button(
            label=f'⬇️ Download Response Letter ({ext.upper()})',
            data=session_state['response_letter_bytes'],
            file_name=f'RTI_Response_{safe_case_id}.{ext}',
            mime=mime_type
        )
        
        from audit_logger import log_response_letter_generated
        log_response_letter_generated(
            case_id=case_id,
            letter_text=f'[Bilingual {ext.upper()} response letter generated]',
            export_format=ext
        )
